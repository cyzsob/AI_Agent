# ingest.py — 文档增量同步入库：扫描 data/documents/ 目录，解析 TXT/PDF/DOCX 文件并存入向量数据库
# 以文件内容 sha256 为指纹，幂等增量同步：
#   - 新增文件 → 解析/分块/嵌入/入库 + 注册
#   - 修改文件 → 删除该文件旧 chunks 后重新入库，库中只保留最新版本
#   - 删除文件（--prune）→ 清理对应 chunks 与注册表记录

import os
import sys
import uuid
import hashlib
import argparse
from pathlib import Path
from dotenv import load_dotenv

load_dotenv()

import psycopg
from langchain_ollama import OllamaEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langchain_postgres import PGVector
from parse_pdf import parse_pdf

DOCS_DIR = Path("data/documents")
SUPPORTED_EXTS = {".txt", ".pdf", ".docx"}
COLLECTION_NAME = "documents"

# ========== 文件解析 ==========


def parse_file(file_path: Path) -> str:
    """根据文件扩展名解析文档内容"""
    file_path = Path(file_path)  # 兼容 str / Path 两种传入方式
    ext = file_path.suffix.lower()

    if ext == ".txt":
        return file_path.read_text(encoding="utf-8")

    if ext == ".pdf":
        return parse_pdf(str(file_path))

    if ext == ".docx":
        from docx import Document as DocxDocument
        doc = DocxDocument(str(file_path))
        # 提取段落 + 表格单元格文本（对齐 mammoth.extractRawText 行为）
        parts = [p.text for p in doc.paragraphs if p.text]
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text for cell in row.cells]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)

    raise ValueError(f"不支持的文件格式: {ext}")


# ========== 文件指纹 ==========


def compute_file_hash(file_path: Path) -> str:
    """计算文件内容 sha256（64KB 分块读取，节省内存）"""
    h = hashlib.sha256()
    with open(file_path, "rb") as f:
        for block in iter(lambda: f.read(65536), b""):
            h.update(block)
    return h.hexdigest()


# ========== 注册表表 doc_file_registry（记录文件指纹，用于增量判断） ==========

_REGISTRY_DDL = """
CREATE TABLE IF NOT EXISTS doc_file_registry (
    file_path   TEXT PRIMARY KEY,
    file_hash   TEXT NOT NULL,
    file_size   BIGINT NOT NULL,
    chunk_count INT  NOT NULL DEFAULT 0,
    updated_at  TIMESTAMPTZ NOT NULL DEFAULT now()
)
"""


def _db_config() -> dict:
    return {
        "host": os.getenv("DB_HOST", "localhost"),
        "port": int(os.getenv("DB_PORT", "5432")),
        "user": os.getenv("DB_USER", "postgres"),
        "password": os.getenv("DB_PASSWORD", ""),
        "dbname": os.getenv("DB_NAME", "RAG_test"),
    }


def _open_conn():
    return psycopg.connect(**_db_config())


def ensure_registry_table(conn):
    conn.execute(_REGISTRY_DDL)
    conn.commit()


def load_registry(conn) -> dict:
    """返回 {file_path: {"hash": ...}}"""
    rows = conn.execute(
        "SELECT file_path, file_hash FROM doc_file_registry"
    ).fetchall()
    return {row[0]: {"hash": row[1]} for row in rows}


def upsert_registry(conn, file_path: str, file_hash: str, file_size: int, chunk_count: int):
    conn.execute(
        """
        INSERT INTO doc_file_registry (file_path, file_hash, file_size, chunk_count, updated_at)
        VALUES (%s, %s, %s, %s, now())
        ON CONFLICT (file_path) DO UPDATE SET
            file_hash   = EXCLUDED.file_hash,
            file_size   = EXCLUDED.file_size,
            chunk_count = EXCLUDED.chunk_count,
            updated_at  = now()
        """,
        (file_path, file_hash, file_size, chunk_count),
    )


def delete_registry_row(conn, file_path: str):
    conn.execute("DELETE FROM doc_file_registry WHERE file_path = %s", (file_path,))


# ========== chunk 删除（与 langchain_postgres 标准 schema 对齐） ==========


def delete_chunks_by_source(conn, source: str, collection_name: str = COLLECTION_NAME):
    """删除指定 source（相对路径）的全部 chunks"""
    conn.execute(
        """
        DELETE FROM langchain_pg_embedding e
        USING langchain_pg_collection c
        WHERE e.collection_id = c.uuid
          AND c.name = %s
          AND e.cmetadata->>'source' = %s
        """,
        (collection_name, source),
    )


def delete_all_chunks(conn, collection_name: str = COLLECTION_NAME):
    """清空指定 collection 的全部 chunks"""
    conn.execute(
        """
        DELETE FROM langchain_pg_embedding e
        USING langchain_pg_collection c
        WHERE e.collection_id = c.uuid AND c.name = %s
        """,
        (collection_name,),
    )


# ========== 增量同步主流程 ==========


def sync_documents(*, prune: bool = False, rebuild: bool = False) -> dict:
    """增量同步 data/documents/ 目录到向量库（幂等）。

    Args:
        prune: 是否清理磁盘上已删除文档对应的 chunks 与注册表记录
        rebuild: 是否先清空向量库再全量重建

    Returns:
        dict: {"added": n, "updated": n, "skipped": n, "removed": n, "errors": [...]}
    """
    result = {"added": 0, "updated": 0, "skipped": 0, "removed": 0, "errors": []}

    # ---------- 1. 目录检查 ----------
    if not DOCS_DIR.is_dir():
        raise RuntimeError(f"无法读取目录 {DOCS_DIR}，请确认目录存在。")

    # ---------- 2. 扫描目录（复用现有扩展名过滤） ----------
    files = [f for f in DOCS_DIR.iterdir() if f.is_file()]
    target_files = sorted(f for f in files if f.suffix.lower() in SUPPORTED_EXTS)

    if not target_files:
        print(f"未在 {DOCS_DIR}/ 目录下找到支持的文档文件。")
        print(f"支持的文件格式: {', '.join(sorted(SUPPORTED_EXTS))}")

    # ---------- 3. 计算指纹，与注册表对比，确定待处理文件 ----------
    conn = _open_conn()
    try:
        ensure_registry_table(conn)

        if rebuild:
            delete_all_chunks(conn)
            delete_registry_all(conn)
            conn.commit()
            print("已清空向量库与注册表，执行全量重建...")

        registry = load_registry(conn)

        # 3.1 磁盘 → 注册表对比
        files_to_index = []  # (rel_path, file_path, sha, is_new)
        disk_paths = set()
        for f in target_files:
            sha = compute_file_hash(f)
            rel = str(f.relative_to(DOCS_DIR))
            disk_paths.add(rel)
            reg = registry.get(rel)
            if reg and reg["hash"] == sha:
                result["skipped"] += 1
                print(f"  - 跳过(未变化): {rel}")
            else:
                is_new = reg is None
                files_to_index.append((rel, f, sha, is_new))
                print(f"  - {'新增' if is_new else '更新'}: {rel}")

        # 3.2 注册表 → 磁盘对比（仅 --prune 时清理已删除文档）
        if prune:
            for rel in list(registry):
                if rel not in disk_paths:
                    delete_chunks_by_source(conn, rel)
                    delete_registry_row(conn, rel)
                    result["removed"] += 1
                    print(f"  - 移除(磁盘已删除): {rel}")
            if result["removed"]:
                conn.commit()

        # ---------- 4. 无变更时直接返回（避免无谓初始化嵌入模型/向量库） ----------
        if not files_to_index:
            print("所有文档均已是最新，无需处理。")
            return result

        # ---------- 5. 初始化嵌入模型 / 向量库 / 分块器 ----------
        embeddings = OllamaEmbeddings(
            model="bge-m3",
            base_url="http://localhost:11434",
        )

        db = _db_config()
        connection_string = (
            f"postgresql+psycopg://{db['user']}:{db['password']}"
            f"@{db['host']}:{db['port']}/{db['dbname']}"
        )

        vector_store = PGVector(
            embeddings=embeddings,
            collection_name=COLLECTION_NAME,
            connection=connection_string,
            use_jsonb=True,
        )

        splitter = RecursiveCharacterTextSplitter(
            chunk_size=200,
            chunk_overlap=20,
        )

        # ---------- 6. 逐个文件处理：解析 → 分块 → 删旧 → 嵌入入库 → 更新注册表 ----------
        for rel, file_path, sha, is_new in files_to_index:
            try:
                print(f"正在解析: {rel}")
                content = parse_file(file_path)
                if not content or not content.strip():
                    print(f"  ⚠ 文件 {rel} 解析结果为空，跳过")
                    result["errors"].append(f"{rel}: 解析结果为空")
                    continue

                raw_doc = Document(page_content=content, metadata={"source": rel})
                chunks = splitter.split_documents([raw_doc])
                if not chunks:
                    print(f"  ⚠ 文件 {rel} 分块结果为空，跳过")
                    result["errors"].append(f"{rel}: 分块结果为空")
                    continue

                # 先删除该文件的旧 chunks（新增时无旧数据，无副作用）
                delete_chunks_by_source(conn, rel)
                conn.commit()

                # 写入增强元数据（便于诊断与按文件清理）
                total = len(chunks)
                for i, chunk in enumerate(chunks):
                    chunk.metadata.update({
                        "file_hash": sha,
                        "chunk_index": i,
                        "total_chunks": total,
                    })

                ids = [str(uuid.uuid4()) for _ in chunks]
                vector_store.add_documents(chunks, ids=ids)

                upsert_registry(conn, rel, sha, file_path.stat().st_size, total)
                conn.commit()

                result["added" if is_new else "updated"] += 1
                print(f"  ✓ {rel}: {total} 个文档块已{'新增' if is_new else '更新'} ({len(content)} 字符)")
            except Exception as err:
                # 处理失败保留旧数据与注册表不动
                result["errors"].append(f"{rel}: {err}")
                print(f"  ✗ 处理失败: {rel} - {err}")
    finally:
        conn.close()

    return result


def delete_registry_all(conn):
    conn.execute("DELETE FROM doc_file_registry")


# ========== CLI 入口 ==========


def main():
    parser = argparse.ArgumentParser(description="增量同步 data/documents/ 目录到向量库（幂等）")
    parser.add_argument(
        "--rebuild", action="store_true",
        help="先清空向量库与注册表，再全量重建（schema/嵌入模型变更时使用）",
    )
    parser.add_argument(
        "--prune", action="store_true",
        help="清理磁盘上已删除文档对应的 chunks 与注册表记录（默认关闭，避免误删）",
    )
    args = parser.parse_args()

    try:
        result = sync_documents(prune=args.prune, rebuild=args.rebuild)
    except Exception as err:
        print(f"❌ 同步失败: {err}")
        sys.exit(1)

    print("\n========== 同步完成 ==========")
    print(
        f"新增: {result['added']} | 更新: {result['updated']} | "
        f"跳过: {result['skipped']} | 移除: {result['removed']}"
    )
    if result["errors"]:
        print(f"错误 ({len(result['errors'])}):")
        for err in result["errors"]:
            print(f"  - {err}")
    print("提示：检索服务运行时可通过 POST /api/admin/sync 触发同步，无需重启。")


if __name__ == "__main__":
    main()
